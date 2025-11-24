"""Utility functions for MEDDPICC coach.

Document parsing and content extraction utilities.
"""

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


async def parse_pdf_bytes(pdf_bytes: bytes) -> str:
    """Parse PDF content from bytes.

    Args:
        pdf_bytes: PDF file as bytes

    Returns:
        Extracted text content
    """
    try:
        import PyPDF2

        pdf_file = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            text_parts.append(f"--- Page {page_num} ---\n{text}\n")

        full_text = "\n".join(text_parts)
        logger.info(
            f"Parsed PDF: {len(pdf_reader.pages)} pages, {len(full_text)} chars"
        )
        return full_text

    except ImportError:
        logger.warning("PyPDF2 not installed - cannot parse PDF")
        return ""
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        return ""


async def parse_docx_bytes(docx_bytes: bytes) -> str:
    """Parse DOCX content from bytes.

    Args:
        docx_bytes: DOCX file as bytes

    Returns:
        Extracted text content
    """
    try:
        import docx

        docx_file = io.BytesIO(docx_bytes)
        doc = docx.Document(docx_file)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        full_text = "\n\n".join(text_parts)
        logger.info(
            f"Parsed DOCX: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars"
        )
        return full_text

    except ImportError:
        logger.warning("python-docx not installed - cannot parse DOCX")
        return ""
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        return ""


async def parse_document(file_content: bytes, filename: str) -> tuple[str, str]:
    """Parse document content based on file extension.

    Args:
        file_content: File content as bytes
        filename: Original filename (used to determine type)

    Returns:
        Tuple of (parsed_content, file_type)
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        content = await parse_pdf_bytes(file_content)
        return content, "PDF"
    elif filename_lower.endswith(".docx"):
        content = await parse_docx_bytes(file_content)
        return content, "DOCX"
    elif filename_lower.endswith(".txt"):
        try:
            content = file_content.decode("utf-8")
            return content, "TXT"
        except Exception as e:
            logger.error(f"Text file decode error: {e}")
            return "", "TXT"
    else:
        logger.warning(f"Unsupported file type: {filename}")
        return "", "UNKNOWN"


async def download_slack_file(file_url: str, token: str) -> bytes | None:
    """Download file from Slack.

    Args:
        file_url: Slack file URL
        token: Slack bot token

    Returns:
        File content as bytes, or None if download fails
    """
    try:
        import aiohttp

        headers = {"Authorization": f"Bearer {token}"}

        async with (
            aiohttp.ClientSession() as session,
            session.get(file_url, headers=headers) as response,
        ):
            if response.status == 200:
                file_bytes = await response.read()
                logger.info(f"Downloaded {len(file_bytes)} bytes from Slack")
                return file_bytes
            else:
                logger.error(f"Slack file download failed: {response.status}")
                return None

    except Exception as e:
        logger.error(f"Error downloading Slack file: {e}")
        return None


async def process_slack_files(files: list[dict[str, Any]], token: str) -> str:
    """Process multiple Slack file attachments.

    Args:
        files: List of Slack file objects with url_private, name
        token: Slack bot token

    Returns:
        Combined parsed content from all files
    """
    parsed_parts = []

    for file_info in files:
        filename = file_info.get("name", "unknown")
        file_url = file_info.get("url_private")

        if not file_url:
            logger.warning(f"No URL for file: {filename}")
            continue

        logger.info(f"Processing file: {filename}")

        # Download file
        file_bytes = await download_slack_file(file_url, token)
        if not file_bytes:
            continue

        # Parse based on file type
        content, file_type = await parse_document(file_bytes, filename)

        if content:
            parsed_parts.append(
                f"## Attachment: {filename} ({file_type})\n\n{content}\n\n"
            )
            logger.info(f"Parsed {file_type} file: {filename} ({len(content)} chars)")
        else:
            logger.warning(f"Could not parse file: {filename}")

    combined_content = "\n---\n\n".join(parsed_parts) if parsed_parts else ""
    return combined_content
